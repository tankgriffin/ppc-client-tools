#!/usr/bin/env python3
"""
Simple Document Exporter for Testing Frameworks
Works without complex dependencies - creates Word docs and beautiful HTML
"""

import os
import re
from pathlib import Path
from datetime import datetime
import argparse

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class SimpleDocumentExporter:
    def __init__(self, client_name: str, project_path: str = None):
        self.client_name = client_name
        self.project_path = Path(project_path) if project_path else Path(f"./{client_name}")
        self.output_dir = self.project_path / "exports"
        self.output_dir.mkdir(exist_ok=True)
        
    def export_to_word(self, markdown_content: str, output_filename: str = None) -> str:
        """Convert markdown to professional Word document"""
        if not DOCX_AVAILABLE:
            print("❌ Word export requires python-docx. Install with: pip install python-docx")
            return None
        
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            output_filename = f"{self.client_name}_Testing_Framework_{timestamp}.docx"
        
        output_path = self.output_dir / output_filename
        
        # Create Word document
        doc = Document()
        
        # Set up document properties
        doc.core_properties.title = f"{self.client_name} - Testing Framework"
        doc.core_properties.author = "PPC Campaign Planning System"
        doc.core_properties.subject = "6-Month Testing Framework"
        
        # Add custom styles
        self._add_word_styles(doc)
        
        # Parse markdown and convert to Word
        self._convert_markdown_to_word(markdown_content, doc)
        
        # Save document
        doc.save(output_path)
        
        print(f"✅ Word document saved: {output_path}")
        return str(output_path)
    
    def export_to_html(self, markdown_content: str, output_filename: str = None) -> str:
        """Convert markdown to beautiful HTML (Google Docs friendly)"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            output_filename = f"{self.client_name}_Testing_Framework_{timestamp}.html"
        
        output_path = self.output_dir / output_filename
        
        # Convert markdown to HTML
        html_content = self._markdown_to_html(markdown_content)
        
        # Create full HTML document
        full_html = self._create_html_document(html_content)
        
        # Save HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        print(f"✅ HTML document saved: {output_path}")
        print(f"💡 Open in browser and copy/paste to Google Docs for perfect formatting")
        return str(output_path)
    
    def export_both_formats(self, markdown_content: str) -> dict:
        """Export to both Word and HTML"""
        results = {}
        
        # Export to HTML (always works)
        try:
            html_path = self.export_to_html(markdown_content)
            results['html'] = html_path
        except Exception as e:
            results['html_error'] = str(e)
        
        # Export to Word (if available)
        try:
            if DOCX_AVAILABLE:
                word_path = self.export_to_word(markdown_content)
                results['word'] = word_path
            else:
                results['word_info'] = "Install python-docx for Word export: pip install python-docx"
        except Exception as e:
            results['word_error'] = str(e)
        
        return results
    
    def _add_word_styles(self, doc):
        """Add professional styles to Word document with Montserrat font"""
        styles = doc.styles
        
        # Custom Title style
        try:
            title_style = styles.add_style('Framework Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Montserrat'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = None  # Use default color
            
            title_para = title_style.paragraph_format
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.space_after = Pt(12)
        except ValueError:
            pass  # Style might already exist
        
        # Custom Heading styles
        try:
            h1_style = styles.add_style('Framework H1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Montserrat'
            h1_font.size = Pt(18)
            h1_font.bold = True
            
            h1_para = h1_style.paragraph_format
            h1_para.space_before = Pt(12)
            h1_para.space_after = Pt(6)
        except ValueError:
            pass
        
        try:
            h2_style = styles.add_style('Framework H2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Montserrat'
            h2_font.size = Pt(14)
            h2_font.bold = True
            
            h2_para = h2_style.paragraph_format
            h2_para.space_before = Pt(8)
            h2_para.space_after = Pt(4)
        except ValueError:
            pass
        
        # Set default paragraph font to Montserrat
        try:
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Montserrat'
            normal_font.size = Pt(11)
        except:
            pass
    
    def _convert_markdown_to_word(self, markdown_content: str, doc):
        """Convert markdown content to Word document elements"""
        lines = markdown_content.split('\n')
        current_table = None
        table_headers = []
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                if current_table is not None:
                    current_table = None
                    table_headers = []
                continue
            
            # Main title (# )
            if line_stripped.startswith('# '):
                title = line_stripped[2:].strip()
                try:
                    paragraph = doc.add_paragraph(title)
                    paragraph.style = 'Framework Title'
                except:
                    paragraph = doc.add_paragraph(title)
                    paragraph.runs[0].bold = True
                    paragraph.runs[0].font.size = Pt(24)
            
            # Heading 1 (## )
            elif line_stripped.startswith('## '):
                heading = line_stripped[3:].strip()
                try:
                    paragraph = doc.add_paragraph(heading)
                    paragraph.style = 'Framework H1'
                except:
                    paragraph = doc.add_paragraph(heading)
                    paragraph.runs[0].bold = True
                    paragraph.runs[0].font.size = Pt(18)
            
            # Heading 2 (### )
            elif line_stripped.startswith('### '):
                heading = line_stripped[4:].strip()
                try:
                    paragraph = doc.add_paragraph(heading)
                    paragraph.style = 'Framework H2'
                except:
                    paragraph = doc.add_paragraph(heading)
                    paragraph.runs[0].bold = True
                    paragraph.runs[0].font.size = Pt(14)
            
            # Heading 3 (#### )
            elif line_stripped.startswith('#### '):
                heading = line_stripped[5:].strip()
                paragraph = doc.add_paragraph(heading)
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)
            
            # Table detection
            elif '|' in line_stripped and line_stripped.startswith('|'):
                if current_table is None:
                    # New table - extract headers
                    headers = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
                    table_headers = headers
                    current_table = doc.add_table(rows=1, cols=len(headers))
                    current_table.style = 'Table Grid'
                    
                    # Add headers
                    for i, header in enumerate(headers):
                        cell = current_table.rows[0].cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True
                elif not line_stripped.startswith('|---'):
                    # Data row
                    cells_data = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
                    if len(cells_data) == len(table_headers):
                        row = current_table.add_row()
                        for i, cell_data in enumerate(cells_data):
                            row.cells[i].text = cell_data
            
            # Bullet points (- )
            elif line_stripped.startswith('- '):
                text = line_stripped[2:].strip()
                paragraph = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(paragraph, text)
                paragraph.style = 'List Bullet'
            
            # Numbered points (1. )
            elif re.match(r'^\d+\.\s', line_stripped):
                text = re.sub(r'^\d+\.\s', '', line_stripped)
                paragraph = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(paragraph, text)
                paragraph.style = 'List Number'
            
            # Bold emphasis (**text**)
            elif '**' in line_stripped:
                paragraph = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(paragraph, line_stripped)
            
            # Regular paragraph
            else:
                if line_stripped:  # Skip empty lines
                    paragraph = doc.add_paragraph()
                    self._add_formatted_text_to_paragraph(paragraph, line_stripped)
    
    def _add_formatted_text_to_paragraph(self, paragraph, text):
        """Add text with bold/italic formatting to paragraph"""
        # Handle bold text (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            else:
                # Regular text
                paragraph.add_run(part)
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML with proper formatting"""
        lines = markdown_content.split('\n')
        html_lines = []
        in_table = False
        table_lines = []
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                if in_table:
                    # End table
                    html_lines.append(self._create_html_table(table_lines))
                    table_lines = []
                    in_table = False
                html_lines.append('<br>')
                continue
            
            # Main title (# )
            if line_stripped.startswith('# '):
                title = line_stripped[2:].strip()
                html_lines.append(f'<h1 class="main-title">{title}</h1>')
            
            # Heading 1 (## )
            elif line_stripped.startswith('## '):
                heading = line_stripped[3:].strip()
                html_lines.append(f'<h2 class="section-heading">{heading}</h2>')
            
            # Heading 2 (### )
            elif line_stripped.startswith('### '):
                heading = line_stripped[4:].strip()
                html_lines.append(f'<h3 class="subsection-heading">{heading}</h3>')
            
            # Heading 3 (#### )
            elif line_stripped.startswith('#### '):
                heading = line_stripped[5:].strip()
                html_lines.append(f'<h4 class="test-heading">{heading}</h4>')
            
            # Table detection
            elif '|' in line_stripped and line_stripped.startswith('|'):
                if not in_table:
                    in_table = True
                table_lines.append(line_stripped)
            
            # Bullet points (- )
            elif line_stripped.startswith('- '):
                if in_table:
                    html_lines.append(self._create_html_table(table_lines))
                    table_lines = []
                    in_table = False
                text = line_stripped[2:].strip()
                formatted_text = self._format_html_text(text)
                html_lines.append(f'<li>{formatted_text}</li>')
            
            # Regular paragraph
            else:
                if in_table:
                    html_lines.append(self._create_html_table(table_lines))
                    table_lines = []
                    in_table = False
                
                formatted_text = self._format_html_text(line_stripped)
                html_lines.append(f'<p>{formatted_text}</p>')
        
        # Handle any remaining table
        if in_table and table_lines:
            html_lines.append(self._create_html_table(table_lines))
        
        return '\n'.join(html_lines)
    
    def _create_html_table(self, table_lines: list) -> str:
        """Create HTML table from markdown table lines"""
        if not table_lines:
            return ''
        
        html = ['<table class="framework-table">']
        
        for i, line in enumerate(table_lines):
            if line.startswith('|---'):
                continue  # Skip separator line
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if i == 0:  # Header row
                html.append('<thead><tr>')
                for cell in cells:
                    html.append(f'<th>{self._format_html_text(cell)}</th>')
                html.append('</tr></thead><tbody>')
            else:  # Data row
                html.append('<tr>')
                for cell in cells:
                    formatted_cell = self._format_html_text(cell)
                    # Add priority styling
                    if 'HIGH' in cell:
                        html.append(f'<td class="priority-high">{formatted_cell}</td>')
                    elif 'MEDIUM' in cell:
                        html.append(f'<td class="priority-medium">{formatted_cell}</td>')
                    elif 'LOW' in cell:
                        html.append(f'<td class="priority-low">{formatted_cell}</td>')
                    else:
                        html.append(f'<td>{formatted_cell}</td>')
                html.append('</tr>')
        
        html.append('</tbody></table>')
        return '\n'.join(html)
    
    def _format_html_text(self, text: str) -> str:
        """Format text with bold and other HTML formatting"""
        # Handle bold text (**text**)
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        return text
    
    def _create_html_document(self, content: str) -> str:
        """Create complete HTML document with professional styling"""
        timestamp = datetime.now().strftime("%B %d, %Y")
        
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.client_name} - Testing Framework</title>
    <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        body {{
            font-family: 'Montserrat', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background-color: #ffffff;
            font-weight: 400;
        }}
        
        .main-title {{
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 15px;
            margin-bottom: 30px;
            font-size: 28px;
            font-weight: 700;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .section-heading {{
            color: #34495e;
            font-size: 22px;
            margin-top: 35px;
            margin-bottom: 15px;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            background-color: #f8f9fa;
            padding: 10px 15px;
            font-weight: 600;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .subsection-heading {{
            color: #2980b9;
            font-size: 18px;
            margin-top: 25px;
            margin-bottom: 12px;
            font-weight: 600;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .test-heading {{
            color: #e74c3c;
            font-size: 16px;
            margin-top: 20px;
            margin-bottom: 10px;
            font-weight: 600;
            background-color: #fff5f5;
            padding: 8px 12px;
            border-left: 3px solid #e74c3c;
            font-family: 'Montserrat', sans-serif;
        }}
        
        p {{
            margin-bottom: 12px;
            text-align: justify;
            font-family: 'Montserrat', sans-serif;
            font-weight: 400;
        }}
        
        ul, ol {{
            margin-bottom: 15px;
            padding-left: 25px;
            font-family: 'Montserrat', sans-serif;
        }}
        
        li {{
            margin-bottom: 6px;
            font-weight: 400;
        }}
        
        .framework-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 25px 0;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            border-radius: 8px;
            overflow: hidden;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .framework-table th {{
            background-color: #3498db;
            color: white;
            padding: 15px 12px;
            text-align: left;
            font-weight: 600;
            font-size: 14px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .framework-table td {{
            border: 1px solid #e0e0e0;
            padding: 12px;
            text-align: left;
            font-size: 14px;
            font-family: 'Montserrat', sans-serif;
            font-weight: 400;
        }}
        
        .framework-table tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}
        
        .framework-table tr:hover {{
            background-color: #e3f2fd;
        }}
        
        .priority-high {{
            color: #e74c3c;
            font-weight: 600;
        }}
        
        .priority-medium {{
            color: #f39c12;
            font-weight: 600;
        }}
        
        .priority-low {{
            color: #27ae60;
            font-weight: 500;
        }}
        
        strong {{
            color: #2c3e50;
            font-weight: 600;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .document-footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #e0e0e0;
            color: #666;
            font-size: 12px;
            text-align: center;
            font-family: 'Montserrat', sans-serif;
            font-weight: 300;
        }}
        
        @media print {{
            body {{
                font-size: 12pt;
                line-height: 1.4;
            }}
            .main-title {{
                font-size: 18pt;
            }}
            .section-heading {{
                font-size: 14pt;
            }}
        }}
        
        /* Google Docs copy-paste optimization */
        * {{
            box-sizing: border-box;
            font-family: 'Montserrat', sans-serif;
        }}
        
        .framework-table {{
            border-spacing: 0;
        }}
    </style>
</head>
<body>
    {content}
    
    <div class="document-footer">
        Generated on {timestamp} by PPC Campaign Planning System
    </div>
</body>
</html>"""

def main():
    """Main function for simple document export"""
    parser = argparse.ArgumentParser(description='Export testing framework to documents (Simple Version)')
    parser.add_argument('client_name', help='Name of the client')
    parser.add_argument('input_file', help='Path to markdown file')
    parser.add_argument('--format', choices=['word', 'html', 'both'], default='both')
    parser.add_argument('--output-name', help='Custom output filename')
    
    args = parser.parse_args()
    
    # Load markdown content
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"❌ Input file not found: {input_path}")
        return
    
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Export documents
    exporter = SimpleDocumentExporter(args.client_name)
    
    print(f"🔄 Converting {args.input_file} for {args.client_name}...")
    
    if args.format == 'both':
        results = exporter.export_both_formats(markdown_content)
        print(f"\n📄 Export Results:")
        for format_type, result in results.items():
            if 'error' in format_type:
                print(f"❌ {format_type}: {result}")
            elif 'info' in format_type:
                print(f"💡 {format_type}: {result}")
            else:
                print(f"✅ {format_type}: {result}")
    
    elif args.format == 'word':
        if DOCX_AVAILABLE:
            exporter.export_to_word(markdown_content, args.output_name)
        else:
            print("❌ Word export requires: pip install python-docx")
    
    elif args.format == 'html':
        exporter.export_to_html(markdown_content, args.output_name)

if __name__ == "__main__":
    main()